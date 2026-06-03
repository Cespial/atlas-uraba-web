# -*- coding: utf-8 -*-
"""
Genera los derechos de peticion (.docx) de la FASE 3 del proyecto Atlas Urabá.
Tensor — gestión formal de datos geográficos ante entidades públicas.
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "/Users/cristianespinal/atlas-uraba-web/docs/peticiones"
os.makedirs(OUT_DIR, exist_ok=True)

# Fecha de hoy (no se obtiene de datos externos; es redacción)
HOY = date(2026, 6, 3)
MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
         "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
FECHA_TXT = f"Medellín, {HOY.day} de {MESES[HOY.month]} de {HOY.year}"

# ----- Paleta / estilo -----
AZUL = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x55, 0x55, 0x55)

PLACEHOLDER = "[VERIFICAR]"

CONTACTO = [
    ("Entidad solicitante", f"Tensor {PLACEHOLDER}"),
    ("NIT", f"{PLACEHOLDER}"),
    ("Representante / Responsable del proyecto", f"{PLACEHOLDER}"),
    ("Cargo", f"{PLACEHOLDER}"),
    ("Dirección de notificación física", f"{PLACEHOLDER}"),
    ("Correo electrónico para notificaciones", f"{PLACEHOLDER}"),
    ("Teléfono de contacto", f"{PLACEHOLDER}"),
]


def set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_membrete(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("TENSOR")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    rs = sub.add_run("Proyecto Atlas Urabá — Inteligencia geoespacial del territorio")
    rs.font.size = Pt(9)
    rs.font.color.rgb = GRIS
    rs.italic = True
    sub2 = doc.add_paragraph()
    rs2 = sub2.add_run(f"NIT {PLACEHOLDER}  ·  {PLACEHOLDER} (correo)  ·  {PLACEHOLDER} (tel.)")
    rs2.font.size = Pt(8)
    rs2.font.color.rgb = GRIS
    # línea divisoria
    line = doc.add_paragraph()
    rl = line.add_run("_" * 95)
    rl.font.color.rgb = AZUL
    rl.font.size = Pt(6)


def add_fecha_y_radicacion(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(FECHA_TXT)
    r.font.size = Pt(11)
    rad = doc.add_paragraph()
    rr = rad.add_run(f"Radicado de salida: {PLACEHOLDER}")
    rr.font.size = Pt(9)
    rr.font.color.rgb = GRIS
    rad.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_destinatario(doc, lineas):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Señores")
    r.font.size = Pt(11)
    for i, ln in enumerate(lineas):
        pl = doc.add_paragraph()
        rl = pl.add_run(ln)
        rl.font.size = Pt(11)
        if i == 0:
            rl.bold = True
            rl.font.color.rgb = AZUL


def add_ref_y_asunto(doc, ref, asunto):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Referencia: ")
    r.bold = True
    p.add_run(ref)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Asunto: ")
    r2.bold = True
    p2.add_run(asunto)


def add_heading(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = AZUL
    return p


def add_parrafo(doc, txt, justify=True):
    p = doc.add_paragraph(txt)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_lista_num(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_firma(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Atentamente,")
    for _ in range(3):
        doc.add_paragraph()
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"{PLACEHOLDER}")
    r2.bold = True
    p3 = doc.add_paragraph()
    p3.add_run(f"Representante / Responsable del proyecto Atlas Urabá — Tensor")
    doc.add_paragraph()
    add_heading(doc, "Datos de contacto para notificación de la respuesta")
    for etiqueta, valor in CONTACTO:
        pc = doc.add_paragraph()
        rc = pc.add_run(f"{etiqueta}: ")
        rc.bold = True
        rc.font.size = Pt(10)
        rv = pc.add_run(valor)
        rv.font.size = Pt(10)


def add_fundamento_legal(doc, extra_normas=None):
    add_heading(doc, "Fundamento legal")
    base = (
        "La presente solicitud se formula en ejercicio del derecho fundamental de petición "
        "consagrado en el artículo 23 de la Constitución Política y reglamentado por la Ley 1755 "
        "de 2015, que sustituyó el Título II de la Ley 1437 de 2011 (Código de Procedimiento "
        "Administrativo y de lo Contencioso Administrativo). De conformidad con la Ley 1712 de "
        "2014 (Ley de Transparencia y del Derecho de Acceso a la Información Pública Nacional), "
        "la información geográfica y catastral en poder de las entidades públicas constituye "
        "información pública sujeta al principio de máxima publicidad (artículo 2), por lo que su "
        "entrega procede de manera gratuita, salvo los costos de reproducción (artículo 26), y en "
        "formato abierto, reutilizable y procesable por medios electrónicos cuando ello sea "
        "posible (artículos 6 y 11)."
    )
    add_parrafo(doc, base)
    if extra_normas:
        add_parrafo(doc, extra_normas)


def add_formato_y_plazo(doc, formato_txt):
    add_heading(doc, "Formato solicitado")
    add_parrafo(doc, formato_txt)
    add_heading(doc, "Plazo legal de respuesta")
    add_parrafo(
        doc,
        "Solicitamos respetuosamente dar respuesta de fondo a esta petición dentro de los QUINCE "
        "(15) DÍAS HÁBILES siguientes a su recepción, término previsto para las peticiones de "
        "acceso a información pública en el artículo 14 de la Ley 1755 de 2015 y en el artículo 26 "
        "de la Ley 1712 de 2014. En caso de que la entidad considere que la información solicitada "
        "reposa total o parcialmente en otra autoridad, agradecemos remitir la petición al "
        "competente dentro de los cinco (5) días hábiles siguientes e informarnos de tal remisión, "
        "conforme al artículo 21 de la Ley 1755 de 2015. Si la información tuviere carácter "
        "reservado, solicitamos motivar la reserva indicando la norma específica que la sustenta "
        "(artículo 28 de la Ley 1712 de 2014)."
    )


# =========================================================================
# OFICIO TIPO 1 — ESTRATIFICACIÓN SOCIOECONÓMICA (a 3 alcaldías)
# =========================================================================
def oficio_estratificacion(municipio, secretaria_lineas, codigo_divipola):
    doc = Document()
    set_base_font(doc)
    for s in doc.sections:
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

    add_membrete(doc)
    add_fecha_y_radicacion(doc)
    add_destinatario(doc, secretaria_lineas)
    add_ref_y_asunto(
        doc,
        f"Derecho de petición de acceso a información pública — Estratificación socioeconómica "
        f"vigente del municipio de {municipio} (DIVIPOLA {codigo_divipola}).",
        "Solicitud del shapefile de estratificación socioeconómica por manzana y por lado de "
        "manzana, junto con los actos administrativos que la adoptan.",
    )
    doc.add_paragraph()
    add_parrafo(
        doc,
        f"Respetados señores de la Secretaría de Planeación del municipio de {municipio}:",
        justify=False,
    )
    add_parrafo(
        doc,
        "En el marco del proyecto Atlas Urabá, una iniciativa de inteligencia geoespacial para la "
        f"caracterización del territorio de la subregión de Urabá, y en condición de peticionarios "
        "interesados en el acceso a información pública, nos dirigimos a ustedes para solicitar la "
        "estratificación socioeconómica vigente del municipio, en formato geográfico abierto.",
    )

    add_fundamento_legal(
        doc,
        extra_normas=(
            "De manera específica, la competencia municipal sobre la estratificación "
            "socioeconómica se fundamenta en la Ley 142 de 1994 (Régimen de Servicios Públicos "
            "Domiciliarios, artículos 101 a 104), que asigna a los alcaldes la adopción, mediante "
            "decreto, de los resultados de la estratificación realizada conforme a las "
            "metodologías del DANE, así como su conservación y actualización. La estratificación "
            "adoptada por acto administrativo es información pública conforme a la Ley 1712 de 2014."
        ),
    )

    add_heading(doc, "Información solicitada")
    add_lista_num(doc, [
        "Capa geográfica (shapefile) de la estratificación socioeconómica urbana vigente del "
        "municipio, desagregada por LADO DE MANZANA, con el código de manzana y el estrato "
        "asignado a cada lado.",
        "En su defecto, o de manera complementaria, la capa geográfica (shapefile) de "
        "estratificación urbana vigente desagregada por MANZANA, con el código de manzana y su "
        "estrato.",
        "Capa geográfica (shapefile) de la estratificación de fincas y viviendas dispersas "
        "rurales vigente, si el municipio la mantiene.",
        "Copia del o los actos administrativos (decretos) vigentes mediante los cuales se adoptó "
        "la estratificación socioeconómica urbana y rural, con indicación de su fecha de "
        "expedición y de la última actualización general realizada.",
        "Diccionario de datos o tabla de atributos que describa los campos de las capas "
        "entregadas (nombre del campo, tipo y significado).",
        "Sistema de referencia de coordenadas (datum y proyección, p. ej. MAGNA-SIRGAS / "
        "EPSG:4686) en que se entrega la información geográfica.",
        "Fecha de corte o vigencia de la información entregada.",
    ])

    add_formato_y_plazo(
        doc,
        "Solicitamos que la información geográfica se entregue preferiblemente en formato "
        "shapefile (.shp con sus archivos asociados .shx, .dbf, .prj y .cpg) comprimido en un "
        "archivo .zip. De no ser posible ese formato, aceptamos GeoPackage (.gpkg), File "
        "Geodatabase (.gdb) o GeoJSON. Para los actos administrativos, solicitamos copia digital "
        "en PDF. La entrega puede realizarse por correo electrónico a la dirección de "
        "notificaciones indicada al final de este oficio o mediante enlace de descarga.",
    )
    add_firma(doc)
    fname = f"01_peticion_estratificacion_{municipio.lower().replace('ó','o').replace('í','i')}.docx"
    path = os.path.join(OUT_DIR, fname)
    doc.save(path)
    return path


# =========================================================================
# OFICIO TIPO 2 — POT/PBOT/EOT USO DEL SUELO (a 3 alcaldías)
# =========================================================================
def oficio_pot(municipio, secretaria_lineas, codigo_divipola):
    doc = Document()
    set_base_font(doc)
    for s in doc.sections:
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

    add_membrete(doc)
    add_fecha_y_radicacion(doc)
    add_destinatario(doc, secretaria_lineas)
    add_ref_y_asunto(
        doc,
        f"Derecho de petición de acceso a información pública — Plan de Ordenamiento Territorial "
        f"(POT/PBOT/EOT) y clasificación del uso del suelo del municipio de {municipio} "
        f"(DIVIPOLA {codigo_divipola}).",
        "Solicitud de la cartografía oficial de uso y clasificación del suelo en formato "
        "shapefile, junto con el acuerdo municipal y el documento técnico de soporte vigentes.",
    )
    doc.add_paragraph()
    add_parrafo(
        doc,
        f"Respetados señores de la Secretaría de Planeación del municipio de {municipio}:",
        justify=False,
    )
    add_parrafo(
        doc,
        "En el marco del proyecto Atlas Urabá, iniciativa de inteligencia geoespacial sobre la "
        "subregión de Urabá, solicitamos respetuosamente la cartografía oficial del ordenamiento "
        "territorial del municipio, en particular las capas de clasificación y uso del suelo, en "
        "formato geográfico abierto y reutilizable.",
    )

    add_fundamento_legal(
        doc,
        extra_normas=(
            "De manera específica, los instrumentos de ordenamiento territorial (POT, PBOT o EOT "
            "según la categoría del municipio) y su cartografía son adoptados mediante acuerdo del "
            "Concejo Municipal conforme a la Ley 388 de 1997 y al Decreto 1077 de 2015 (Decreto "
            "Único Reglamentario del Sector Vivienda, Ciudad y Territorio). La cartografía del POT "
            "es un documento público que hace parte integral del acto administrativo que lo adopta "
            "y, por tanto, es información pública de obligatoria divulgación conforme a la Ley 1712 "
            "de 2014."
        ),
    )

    add_heading(doc, "Información solicitada")
    add_lista_num(doc, [
        "Capa geográfica (shapefile) de CLASIFICACIÓN DEL SUELO del municipio (suelo urbano, de "
        "expansión urbana, rural, suburbano y de protección), conforme al POT/PBOT/EOT vigente.",
        "Capa geográfica (shapefile) de USOS DEL SUELO URBANO (zonificación de usos: residencial, "
        "comercial, industrial, dotacional, mixto, etc.).",
        "Capa geográfica (shapefile) de USOS DEL SUELO RURAL y de las áreas de actividad o "
        "categorías de uso rural definidas por el plan.",
        "Capa geográfica (shapefile) del PERÍMETRO URBANO y del perímetro de expansión vigentes.",
        "Capas geográficas (shapefile) de áreas de protección ambiental, amenazas y riesgos, y "
        "demás capas temáticas del componente cartográfico del POT/PBOT/EOT, si están disponibles.",
        "Copia del Acuerdo Municipal vigente que adopta el POT/PBOT/EOT y de sus modificaciones o "
        "revisiones posteriores, con indicación de su número, fecha y estado de vigencia.",
        "Documento Técnico de Soporte (DTS) y, si existe, la cartografía complementaria publicada.",
        "Diccionario de datos o tabla de atributos de las capas entregadas y el sistema de "
        "referencia de coordenadas (datum y proyección, p. ej. MAGNA-SIRGAS / EPSG:4686).",
        "Indicación del estado del proceso de revisión o ajuste del POT, si actualmente se "
        "encuentra en formulación o revisión.",
    ])

    add_formato_y_plazo(
        doc,
        "Solicitamos que la cartografía se entregue preferiblemente en formato shapefile (.shp con "
        "sus archivos asociados .shx, .dbf, .prj y .cpg) comprimido en un archivo .zip, una capa "
        "por tema. De no ser posible, aceptamos GeoPackage (.gpkg), File Geodatabase (.gdb), DWG "
        "georreferenciado o GeoJSON. El Acuerdo Municipal y el DTS pueden entregarse en PDF. La "
        "entrega puede realizarse por correo electrónico a la dirección de notificaciones "
        "indicada al final de este oficio o mediante enlace de descarga.",
    )
    add_firma(doc)
    fname = f"02_peticion_pot_usosuelo_{municipio.lower().replace('ó','o').replace('í','i')}.docx"
    path = os.path.join(OUT_DIR, fname)
    doc.save(path)
    return path


# =========================================================================
# OFICIO TIPO 3 — ICA: PREDIOS EXPORTADORES + VIGILANCIA Foc TR4
# =========================================================================
def oficio_ica():
    doc = Document()
    set_base_font(doc)
    for s in doc.sections:
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)

    destinatario = [
        "INSTITUTO COLOMBIANO AGROPECUARIO — ICA",
        "Seccional / Gerencia Seccional Urabá",
        "Dirección Técnica de Epidemiología y Vigilancia Fitosanitaria",
        f"Dirección: {PLACEHOLDER} (sede ICA Urabá)",
    ]

    add_membrete(doc)
    add_fecha_y_radicacion(doc)
    add_destinatario(doc, destinatario)
    add_ref_y_asunto(
        doc,
        "Derecho de petición de acceso a información pública — Registro de predios exportadores de "
        "banano y plátano y vigilancia fitosanitaria de Fusarium oxysporum f. sp. cubense raza 4 "
        "tropical (Foc R4T / TR4) en la subregión de Urabá.",
        "Solicitud de los registros y la cartografía de predios productores/exportadores y de los "
        "datos de vigilancia y zonificación del Foc TR4, en formato geográfico abierto.",
    )
    doc.add_paragraph()
    add_parrafo(
        doc,
        "Respetados señores del Instituto Colombiano Agropecuario, seccional Urabá:",
        justify=False,
    )
    add_parrafo(
        doc,
        "En el marco del proyecto Atlas Urabá, iniciativa de inteligencia geoespacial sobre la "
        "subregión de Urabá (Antioquia), y dada la relevancia sanitaria del Fusarium oxysporum "
        "f. sp. cubense raza 4 tropical (Foc R4T / TR4) para la cadena del banano y plátano, "
        "solicitamos respetuosamente la información de registro de predios y de vigilancia "
        "fitosanitaria que reposa en esa entidad, en formato abierto.",
    )

    add_fundamento_legal(
        doc,
        extra_normas=(
            "De manera específica, el ICA es la autoridad sanitaria y fitosanitaria del país "
            "conforme a la Ley 101 de 1993 y al Decreto 4765 de 2008, y administra el registro de "
            "predios productores y exportadores de vegetales y la vigilancia fitosanitaria. La "
            "emergencia sanitaria por Foc R4T fue declarada mediante la Resolución ICA 09444 de "
            "2019 y desarrollada en actos administrativos posteriores sobre medidas de "
            "cuarentena, bioseguridad y registro. La información estadística, los registros y la "
            "cartografía de vigilancia que no tengan reserva legal son información pública "
            "conforme a la Ley 1712 de 2014. Cuando algún dato individual de predios esté "
            "protegido por reserva (p. ej. datos personales del productor conforme a la Ley 1581 "
            "de 2012), solicitamos su entrega de forma anonimizada o agregada, conforme al "
            "principio de divulgación parcial del artículo 21 de la Ley 1712 de 2014."
        ),
    )

    add_heading(doc, "Información solicitada")
    add_lista_num(doc, [
        "Listado y/o capa geográfica (shapefile) de los predios REGISTRADOS COMO EXPORTADORES de "
        "banano y plátano en la jurisdicción de la seccional Urabá, con su ubicación "
        "georreferenciada (punto o polígono), municipio, vereda y número o código de registro "
        "ante el ICA. Si la ubicación predial precisa tiene reserva, solicitamos la información "
        "agregada por vereda o por municipio.",
        "Listado y/o capa geográfica (shapefile) de predios PRODUCTORES de banano y plátano "
        "registrados ante el ICA en la subregión de Urabá, con área sembrada y municipio.",
        "Capa geográfica (shapefile) de la ZONIFICACIÓN de vigilancia y manejo del Foc R4T / TR4 "
        "en Urabá: áreas con presencia confirmada, áreas de contención o cuarentena, zonas "
        "tampón (buffer) y áreas libres, vigentes a la fecha.",
        "Datos de los FOCOS o predios con detección positiva de Foc R4T en Urabá (georreferencia "
        "o ubicación agregada, fecha de detección y estado de manejo), en lo que no tenga reserva.",
        "Estadísticas de vigilancia fitosanitaria del Foc R4T en Urabá: número de predios "
        "monitoreados, muestras tomadas y resultados, por municipio y por periodo disponible.",
        "Relación de los actos administrativos del ICA vigentes sobre el Foc R4T y sobre el "
        "registro de predios exportadores aplicables a Urabá (número, fecha y objeto).",
        "Diccionario de datos o tabla de atributos de las capas entregadas y el sistema de "
        "referencia de coordenadas (datum y proyección, p. ej. MAGNA-SIRGAS / EPSG:4686).",
        "Fecha de corte o vigencia de la información entregada.",
    ])

    add_formato_y_plazo(
        doc,
        "Solicitamos que la información geográfica se entregue preferiblemente en formato "
        "shapefile (.shp con sus archivos asociados .shx, .dbf, .prj y .cpg) comprimido en .zip. "
        "De no ser posible, aceptamos GeoPackage (.gpkg), File Geodatabase (.gdb), KML/KMZ o "
        "GeoJSON. Los listados pueden entregarse en formato tabular abierto (CSV o XLSX) y los "
        "actos administrativos en PDF. La entrega puede realizarse por correo electrónico a la "
        "dirección de notificaciones indicada al final de este oficio o mediante enlace de "
        "descarga.",
    )
    add_firma(doc)
    path = os.path.join(OUT_DIR, "03_peticion_ica_predios_foc_tr4.docx")
    doc.save(path)
    return path


# =========================================================================
# GENERACIÓN
# =========================================================================
municipios = [
    ("Apartadó", [
        "ALCALDÍA MUNICIPAL DE APARTADÓ",
        "Secretaría de Planeación Municipal",
        f"Dirección: {PLACEHOLDER} (Palacio Municipal de Apartadó)",
    ], "05045"),
    ("Turbo", [
        "ALCALDÍA MUNICIPAL DE TURBO",
        "Secretaría de Planeación Municipal",
        f"Dirección: {PLACEHOLDER} (Palacio Municipal de Turbo)",
    ], "05837"),
    ("Chigorodó", [
        "ALCALDÍA MUNICIPAL DE CHIGORODÓ",
        "Secretaría de Planeación Municipal",
        f"Dirección: {PLACEHOLDER} (Palacio Municipal de Chigorodó)",
    ], "05172"),
]

generados = {"estratificacion": [], "pot": [], "ica": None}

for muni, dest, divipola in municipios:
    generados["estratificacion"].append(oficio_estratificacion(muni, dest, divipola))
    generados["pot"].append(oficio_pot(muni, dest, divipola))

generados["ica"] = oficio_ica()

print("GENERADOS:")
for k, v in generados.items():
    print(k, v)
